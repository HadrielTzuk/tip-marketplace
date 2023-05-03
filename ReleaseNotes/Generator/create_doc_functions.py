from __future__ import annotations

import os

import docx
import docx.shared
from docx import section
from docx.text import paragraph

import creator_constants as consts
import settings
from creator_utils import validate_result_folder
from style_datamodels import (
    FooterStyle,
    HeaderStyle,
    MainTextTitleStyle,
    PublishDateStyle,
    TechDetailsTextStyle,
    TechDetailsTitleStyle,
    TitleStyle,
)


def create_doc_rn(
        document: docx.Document,
        release_notes: dict,
        release_version: str,
        results_folder: str,
        minimum_version: str,
        integration_versions: dict[str, float],
        current_year: int,
        formatted_current_date: str,
        header_style: HeaderStyle,
        title_style: TitleStyle,
        publish_time_style: PublishDateStyle,
        tech_details_title_style: TechDetailsTitleStyle,
        tech_details_text_style: TechDetailsTextStyle,
        main_text_style: MainTextTitleStyle,
        footer_style: FooterStyle,
) -> None:
    """
    Create the docx release notes
    Args:
        footer_style (FooterStyle):
        main_text_style (MainTextTitleStyle):
        tech_details_text_style (TechDetailsTextStyle):
        tech_details_title_style (TechDetailsTitleStyle):
        publish_time_style (PublishDateStyle):
        title_style (TitleStyle):
        header_style (HeaderStyle):
        document (docx.Document):
        release_notes (dict): The integration RN collected by collect_release_notes_from_integrations
        release_version (str): The version of the current release
        results_folder (str): The path to the folder of the results of the script
        minimum_version (str): The minimum siemplify version of current release
        integration_versions (dict[str, float]): The integrations' latest versions dict
        current_year (int): The current year
        formatted_current_date (str): A string representing the current date
    """
    validate_result_folder(results_folder)

    _create_header(document, header_style)
    _create_footer(document, footer_style, current_year)
    _create_footer(document, footer_style, current_year, is_first_page=True)

    title_paragraph = _create_title(document, title_style, release_version)
    _add_publish_date_to_title(document, publish_time_style, formatted_current_date, existing_paragraph=title_paragraph)
    _create_tech_details_title(document, tech_details_title_style)
    _create_tech_details_text(document, tech_details_text_style, minimum_version)

    for category in release_notes.keys():
        if release_notes[category]:
            _create_main_text_title(document, main_text_style, category)
            _create_main_text_rn(
                data=release_notes[category],
                parent=consts.ROOT_PARENT,
                document=document,
                category=category,
                integration_versions=integration_versions,
            )

    document.save(os.path.join(results_folder, f'{release_version}{consts.DOCX_EXTENSION}'))


def _create_header(document: docx.Document, header_style: HeaderStyle) -> section.Section:
    """
    Add header to the docx document
    Args:
        document (docx.Document): A docx.Document() object
        header_style (HeaderStyle):
    """
    new_section = document.sections[0]
    new_section.header_distance = header_style.header_distance
    new_section.left_margin = header_style.left_margin
    new_section.right_margin = header_style.right_margin
    new_section.different_first_page_header_footer = header_style.different_first_page_header_footer

    header = new_section.first_page_header
    header.paragraphs[0].add_run().add_picture(consts.HEADER_PIC_PATH, width=new_section.page_width)

    return new_section


def _create_title(
        document: docx.Document,
        title_style: TitleStyle,
        release_version: str,
        existing_paragraph: paragraph.Paragraph | None = None,
) -> paragraph.Paragraph:
    """
    Create the title of the document
    Args:
        document (docx.Document):
        title_style (TitleStyle):
        release_version (str):
    """
    p = existing_paragraph if existing_paragraph is not None else document.add_paragraph()
    p.paragraph_format.alignment = title_style.alignment
    p.paragraph_format.space_before = title_style.space_before

    heading_run = p.add_run(consts.DOC_TITLE.format(release_version))
    heading_run.bold = title_style.is_bold
    heading_run.font.underline = title_style.is_underline
    heading_run.font.name = title_style.font_name
    heading_run.font.size = title_style.font_size
    heading_run.font.color.rgb = title_style.text_color

    return p


def _add_publish_date_to_title(
        document: docx.Document,
        publish_date_style: PublishDateStyle,
        formatted_current_date: str,
        existing_paragraph: paragraph.Paragraph | None = None,
) -> paragraph.Paragraph:
    """
    Adds the publishing time to the title
    Args:
        document (docx.Document):
        publish_date_style (PublishDateStyle):
        formatted_current_date (str):
    """
    p = existing_paragraph if existing_paragraph is not None else document.add_paragraph()
    p.paragraph_format.alignment = publish_date_style.alignment
    p.paragraph_format.space_before = publish_date_style.space_before

    publish_time_run = p.add_run(consts.PUBLISH_TIME.format(formatted_current_date))
    publish_time_run.bold = publish_date_style.is_bold
    publish_time_run.font.underline = publish_date_style.is_underline
    publish_time_run.font.name = publish_date_style.font_name
    publish_time_run.font.size = publish_date_style.font_size
    publish_time_run.font.color.rgb = publish_date_style.text_color

    return p


def _create_tech_details_title(
        document: docx.Document,
        tech_details_style: TechDetailsTitleStyle,
        existing_paragraph: paragraph.Paragraph | None = None,
) -> paragraph.Paragraph:
    """

    Args:
        document (docx.Document):
        tech_details_style (TechDetailsTitleStyle):
        existing_paragraph (paragraph.Paragraph):

    Returns:

    """
    p = existing_paragraph if existing_paragraph is not None else document.add_paragraph()
    p.paragraph_format.alignment = tech_details_style.alignment
    p.paragraph_format.space_before = tech_details_style.space_before
    p.paragraph_format.left_indent = tech_details_style.left_indent

    tech_details_run = p.add_run(consts.TECH_DETAILS_TITLE)
    tech_details_run.bold = tech_details_style.is_bold
    tech_details_run.font.underline = tech_details_style.is_underline
    tech_details_run.font.name = tech_details_style.font_name
    tech_details_run.font.size = tech_details_style.font_size
    tech_details_run.font.color.rgb = tech_details_style.text_color

    return p


def _create_tech_details_text(
        document: docx.Document,
        tech_details_style: TechDetailsTextStyle,
        minimum_version: str,
        existing_paragraph: paragraph.Paragraph | None = None,
) -> paragraph.Paragraph:
    """

    Args:
        existing_paragraph (paragraph.Paragraph):
        document (docx.Document):
        tech_details_style (TechDetailsTextStyle):
        minimum_version (str):

    Returns:

    """
    p = existing_paragraph if existing_paragraph is not None else document.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = tech_details_style.left_indent

    tech_details_text_1 = p.add_run(consts.TECH_DETAILS_TEXT['first part'])
    tech_details_text_1.font.underline = tech_details_style.is_underline
    tech_details_text_1.font.name = tech_details_style.font_name
    tech_details_text_1.font.size = tech_details_style.font_size
    tech_details_text_1.font.color.rgb = tech_details_style.text_color

    tech_details_text_2 = p.add_run(consts.TECH_DETAILS_TEXT['second part'].format(minimum_version))
    tech_details_text_2.bold = True
    tech_details_text_2.font.underline = tech_details_style.is_underline
    tech_details_text_2.font.name = tech_details_style.font_name
    tech_details_text_2.font.size = tech_details_style.font_size
    tech_details_text_2.font.color.rgb = tech_details_style.text_color

    tech_details_text_3 = p.add_run(consts.TECH_DETAILS_TEXT['third part'])
    tech_details_text_3.font.underline = tech_details_style.is_underline
    tech_details_text_3.font.name = tech_details_style.font_name
    tech_details_text_3.font.size = tech_details_style.font_size
    tech_details_text_3.font.color.rgb = tech_details_style.text_color

    return p


def _create_main_text_title(

        document: docx.Document,
        main_text_title_style: MainTextTitleStyle,
        category: str,
        existing_paragraph: paragraph.Paragraph | None = None,
) -> paragraph.Paragraph:
    """

    Args:
        document (docx.Document):
        main_text_title_style (MainTextTitleStyle):
        category (str):
        existing_paragraph (paragraph.Paragraph):

    Returns:

    """
    p = existing_paragraph if existing_paragraph is not None else document.add_paragraph()
    p.paragraph_format.space_before = main_text_title_style.space_before
    p.paragraph_format.space_after = main_text_title_style.space_after
    p.paragraph_format.left_indent = main_text_title_style.left_indent

    main_text = p.add_run(f"{category}?")
    main_text.font.name = main_text_title_style.font_name
    main_text.font.size = main_text_title_style.font_size
    main_text.font.color.rgb = main_text_title_style.text_color
    main_text.bold = main_text_title_style.is_bold
    main_text.font.underline = main_text_title_style.is_underline

    return p


def _create_main_text_rn(
        data: list | dict,
        parent: str,
        document: docx.Document,
        category: str,
        integration_versions: dict[str, float],
) -> None:
    """
    Construct Docx RN from data
    Args:
        data (list | dict):
        parent (str):
        document (Document):
        category (str):
        integration_versions (dict[str, float]): The integration_versions.json file content
    """
    already_exists = list()
    if len(data):
        if isinstance(data, list):
            for item in data:
                p = document.add_paragraph(style='List Bullet 3')
                p.paragraph_format.left_indent = docx.shared.Cm(5.5)
                r = p.add_run(f"[{item.get('ticket')}] {item.get('description')}")
                r.font.size = docx.shared.Pt(11)
                r.font.name = settings.DOCUMENT_FONT
                r.bold = False

        else:
            for k, v in data.items():
                if parent == "root":
                    # An integration
                    p = document.add_paragraph(style='List Bullet')
                    p.paragraph_format.left_indent = docx.shared.Cm(3.5)
                    r = p.add_run(k)
                    r.font.name = settings.DOCUMENT_FONT
                    r.font.size = docx.shared.Pt(12)
                    r.bold = True
                    line = p.add_run(f"\nIntegration version: {integration_versions[k]}")
                    line.font.name = settings.DOCUMENT_FONT
                    line.font.size = docx.shared.Pt(10)
                    line.bold = False

                else:
                    # Item (action/job/connector)
                    if k not in already_exists:
                        generated_text = k
                        if category is not None:
                            if category == "What's New":
                                generated_text = f"New {k}(s):"

                            elif category == "What's Improved":
                                generated_text = f"{k} Update(s):"

                            elif category == "What's Removed":
                                generated_text = f"Removed {k}(s)"

                            elif category == "What's Deprecated":
                                generated_text = f"Deprecated {k}(s)"

                            elif category == "What's Regressed":
                                generated_text = f"Regressed {k}(s)"

                        p = document.add_paragraph(style='List Bullet 2')
                        r = p.add_run(generated_text)
                        p.paragraph_format.left_indent = docx.shared.Cm(4.5)
                        r.font.size = docx.shared.Pt(11)
                        r.font.name = settings.DOCUMENT_FONT
                        r.bold = True
                        already_exists.append(k)
                _create_main_text_rn(v, k, document, category, integration_versions)


def _create_footer(
        document: docx.Document,
        footer_style: FooterStyle,
        current_year: int,
        is_first_page: bool = False,
        existing_paragraph: paragraph.Paragraph | None = None,
) -> paragraph.Paragraph:
    """
    Add footer to the docx document
    Args:
        existing_paragraph (paragraph.Paragraph):
        is_first_page (bool):
        document (docx.Document): A docx.Document().section[0].footer or docx.Document().section[0].first_page_footer
        current_year (int): The current_year
        footer_style (FooterStyle):

    Returns:

    """
    new_section = document.sections[0]
    footer = new_section.first_page_footer if is_first_page else new_section.footer

    table = footer.add_table(
        rows=footer_style.table_rows,
        cols=footer_style.table_columns,
        width=new_section.page_width,
    )
    hdr_cells = table.rows[0].cells
    p = existing_paragraph if existing_paragraph is not None else hdr_cells[0].add_paragraph()
    p.paragraph_format.left_indent = footer_style.left_indent
    p.paragraph_format.space_before = footer_style.space_before

    footer_run = p.add_run(consts.FOOTER_TEXT.format(current_year))
    footer_run.font.name = footer_style.font_name
    footer_run.font.size = footer_style.font_size

    p = hdr_cells[1].add_paragraph()
    p.paragraph_format.alignment = footer_style.alignment
    p.paragraph_format.right_indent = footer_style.right_indent
    p.add_run().add_picture(consts.LOGO_PATH, width=footer_style.logo_width)

    return p
